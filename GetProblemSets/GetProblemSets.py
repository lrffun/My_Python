# coding=utf-8
from lxml import etree
import docx
import requests
import re
import html
import time

#courseId = "81840553"
url = "https://mooc1.chaoxing.com/course/{{courseId}}.html"
mHeaders = {
    'User-Agent': r'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) '
                  r'Chrome/45.0.2454.85 Safari/537.36 115Browser/6.0.3'
}

def getHtml():  #得到页面
    mUrl = url.replace('{{courseId}}',courseId)
    response = requests.get(mUrl,headers=mHeaders)
    response.encoding = 'utf-8'
    mHtml = html.unescape(response.text)        #Html反转义
    #response.encoding = 'utf-8'
    print(mHtml)
    return mHtml

def getCourseUrlList(zhtml):  #得到可以用的每一个课时的Url
    divList = []
    re_rule = 'courseId=' + courseId + '&knowledgeId=(.*?)">'
    # for i in re.findall(re_rule,html):
    #     divList.append(i)
    divList = re.findall(re_rule,zhtml)
    urlList = []
    for i in divList:
        mUrl = 'https://mooc1.chaoxing.com/nodedetailcontroller/visitnodedetail?courseId='+courseId+'&knowledgeId='+i
        print(mUrl)
        try:
            response = requests.get(mUrl,headers=mHeaders,timeout=1)
            if response.status_code == 200:
                if courseId in response.text:
                     urlList.append(mUrl)
                     print('访问成功')
                else:
                    print('非课程网页')
        except Exception as e:
            print('访问失败')

    return urlList

def getZuoYeUrl(urlList):       #得到测验的Url
    tUrlList = []
    for i in urlList:
        response = requests.get(i,headers=mHeaders).content.decode('utf-8')
        res = re.findall('workid&quot;:&quot;(.*?)&quot;,', response)
        if len(res):
            for i in res:
                tUrl = "https://mooc1.chaoxing.com/api/selectWorkQuestion?workId="+i+"&ut=null&classId=0&courseId="+courseId
                tUrlList.append(tUrl)
    return tUrlList

def writeDocx(urlList):          #从测验Url中读取题目并写入Word文档
    #url = "https://mooc1.chaoxing.com/api/selectWorkQuestion?workId=3ba0d4c8e1e740e5a5507487b895c15e&ut=null&classId=0&courseId=81840553"
    for url in urlList:
        mHtml = requests.get(url, headers=mHeaders).content.decode("utf-8")
        file = docx.Document()
        h3 = re.findall('<h3>(.*?)</h3>', mHtml)
        Title = ""
        for i in h3:
            Title = html.unescape(i)
            file.add_heading(Title)

        text = html.unescape(mHtml)
        mHtml = etree.HTML(text)  # 将html转换为xml
        timuList = mHtml.xpath('//div[@class="TiMu"]')  # 找到每一个题目及其所有选项
        for i in timuList:
            time.sleep(0.05)
            mStr = etree.tostring(i).decode('utf-8')  # 将xml树结点读出并转换为utf-8格式
            res = html.unescape(mStr)  # 解码xml
            tType = re.findall('(【.*?】)', res)
            tRType = []
            for a in tType:
                p_rule = '<.*?>'
                tRType.append(re.sub(p_rule,'',str(a)))
            tGan = re.findall('】<?p?>?(.*?)</p>', res)
            if not len(tGan):
                tGan = re.findall('<div class="Zy_TItle_p">(.*?)</div>', res)
            if not len(tGan):
                tGan = re.findall('】(.*?)</div>',res)
            tRGan = []
            for a in tGan:
                p_rule = '<.*?>'
                tRGan.append(re.sub(p_rule,'',str(a)))
            file.add_paragraph(tRType + tRGan)
            '''
                for j in tType:
                print(j)
                file.add_paragraph(j)
            for j in tGan:
                print(j)
                file.add_paragraph(j)
            '''
            XuanXiang = etree.HTML(res)
            tAny = XuanXiang.xpath('//li[@class="clearfix"]')
            for j in tAny:
                tStr = etree.tostring(j).decode('utf-8')
                tRes = html.unescape(tStr)
                tXuan = re.findall('<i class="fl">(.*?)</i>.*?none;"><?p?>?(.*?)<?/?p?>?</a></li>', tRes)
                tRXuan = []
                for a in tXuan:
                    tRRXuan = ""
                    for b in a:
                        p_rule = '<.*?>'
                        tRRXuan = tRRXuan + re.sub(p_rule, '', str(b))
                    tRXuan.append(tRRXuan)
                for k in tRXuan:
                    file.add_paragraph(k)

        file.save("C:\\Users\\HP\\Desktop\\马克思主义原理\\"+Title+".docx")
        print(Title+'爬取完成')
        time.sleep(0.3)

if __name__ == "__main__":
    courseId = "208255733"
    zHtml = getHtml()
    canUseUrl = getCourseUrlList(zHtml)
    zuoYeUrl = getZuoYeUrl(canUseUrl)
    writeDocx(zuoYeUrl)

